"""
Tests for Tier-2 deterministic blank detection (FR-5.1, FR-5.2).

The .docx fixture is built programmatically (no client data is committed) and
exercises every detector: highlighted runs, underscore fill-lines, [ ]/[X]
checkboxes, XXX sentinels, and empty table grids. A regression test pins the
shipped sample templates so normalization never inflates their blank counts.
"""
from __future__ import annotations

import os

from app import blank_detect, templates


def _keys(text: str):
    norm = blank_detect.normalize(text)
    return [f.key for f in templates.detect_fields(norm)]


# --------------------------------------------------------------------------- #
# Text-level conventions
# --------------------------------------------------------------------------- #
def test_underscore_fill_line_named_from_label():
    keys = _keys("NAME OF DECEASED CHILD: ____________________.")
    assert keys == ["name_of_deceased_child"]


def test_wide_bracket_span_and_checkboxes_are_distinguished():
    keys = _keys("DATED:[           ]\n[X] BY MAIL   [  ] BY FACSIMILE")
    # one wide-bracket fill span + two checkboxes = three blanks
    assert len(keys) == 3
    assert keys[0] == "dated"


def test_xrun_sentinels():
    keys = _keys("Email XXX@firm.com regarding claim XXXX.")
    assert "email" in keys
    assert any(k.startswith("xxxx") for k in keys)


def test_existing_markup_is_preserved_and_not_double_counted():
    # An underscore signature line directly above a {{markup}} field is the same
    # field already declared — it must be suppressed, and markup left intact.
    text = "Acknowledged and agreed:\n____________________\n{{client_name}}"
    norm = blank_detect.normalize(text)
    assert "{{client_name}}" in norm
    assert _keys(text) == ["client_name"]


def test_bare_label_colon_is_not_a_blank():
    # Headings like these must not be flagged on their own.
    assert _keys("Sincerely:\nDear Counsel:\nAcknowledged and agreed:") == []


# --------------------------------------------------------------------------- #
# .docx OOXML conventions (highlight + empty table grid)
# --------------------------------------------------------------------------- #
def _build_docx(path: str) -> None:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    d = Document()
    p = d.add_paragraph()
    p.add_run("Claim No. : ")
    hot = p.add_run("1E01E0143810")
    hot.font.highlight_color = WD_COLOR_INDEX.YELLOW
    d.add_paragraph("Affiant residing at ____________________, being duly sworn.")
    d.add_paragraph("[X] BY MAIL   [  ] BY FACSIMILE")
    d.add_paragraph("Please contact XXX@firm.com regarding XXXX.")
    t = d.add_table(rows=3, cols=2)  # header row + two empty body rows
    t.rows[0].cells[0].text = "Name of Child"
    t.rows[0].cells[1].text = "Age"
    d.save(path)


def test_docx_highlight_underscore_checkbox_xrun_table(tmp_path):
    path = os.path.join(tmp_path, "fixture.docx")
    _build_docx(path)
    kind, text = templates.read_template_text(path)
    keys = [f.key for f in templates.detect_fields(text)]

    assert kind == "docx"
    assert "claim_no" in keys            # highlighted run named by its label
    assert "name_of_child" in keys       # empty table column
    assert "age" in keys                 # empty table column
    assert "email" in keys               # XXX@firm.com sentinel
    # underscore fill-line + two checkboxes are present too
    assert len(keys) >= 7


# --------------------------------------------------------------------------- #
# Regression: shipped sample templates keep their counts
# --------------------------------------------------------------------------- #
def test_sample_templates_unchanged():
    root = os.path.join(os.path.dirname(__file__), os.pardir, "data", "templates")
    expected = {
        "demand_letter.md": 14,
        "engagement_letter.md": 9,
        "probate_petition.txt": 9,
    }
    for fname, count in expected.items():
        _, text = templates.read_template_text(os.path.join(root, fname))
        assert len(templates.detect_fields(text)) == count, fname
