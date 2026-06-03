"""
Tier-2 deterministic blank detection (FR-5.1, FR-5.2).

Real firm templates do not use Verbatim's {{key}} / [[key]] markup — they mark
blanks the way lawyers always have: underscore fill-lines, bracketed spans,
[  ] checkboxes, X-run sentinels (XXX / XXXX), "Label :" lines, empty table
grids, and — in .docx — yellow-highlighted runs. This module recognizes those
conventions and *normalizes* them into the canonical {{key | instruction}}
markup the rest of the pipeline already understands, so detection, filling, and
provenance work unchanged.

Detection is deterministic and high-precision by design; the ambiguous semantic
cases (free-form ALL-CAPS, inline prose instructions) are deferred to a later
LLM tier. The full taxonomy, evidence, and rationale are in
docs/blank-detection.md. Normalization runs once, when a template is read.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from typing import List, Optional, Tuple

# Private-use sentinels the .docx reader emits to carry an OOXML signal
# (a yellow-highlighted run) through to the text-level normalizer.
HL_OPEN, HL_CLOSE = "", ""

# Canonical markup already understood by the pipeline: {{k}} / [[k]] / {{k | i}}.
# Pre-existing markup is protected — Tier-2 detectors never fire inside it.
_MARKUP_RE = re.compile(r"(?:\{\{|\[\[)\s*[^|}\]]+?\s*(?:\|[^}\]]*?)?\s*(?:\}\}|\]\])")

# --- Tier-2 text detectors (high precision) -------------------------------- #
_UNDERSCORE_RE = re.compile(r"_{3,}")              # ____ fill lines
_BRACKET_SPAN_RE = re.compile(r"\[[ \t]{4,}\]")    # [        ] wide fill spans
_CHECKBOX_RE = re.compile(r"\[[ \t]{0,3}[xX]?[ \t]{0,3}\]")  # [ ] [x] [X] checkboxes
_XRUN_RE = re.compile(r"(?<![A-Za-z])X{3,}(?:[\w@-]|\.(?=\w))*")  # XXX / XXXX / XXX@host

# A bare "Label :" line (empty value) is too ambiguous to flag as a blank on its
# own — it also matches headings like "Acknowledged and agreed:" or "Sincerely:".
# The useful part of that signal (using a preceding label to *name* an adjacent
# blank) lives in _label_before(), not as a standalone detector.
_TEXT_DETECTORS = [
    ("checkbox", _CHECKBOX_RE),
    ("bracket", _BRACKET_SPAN_RE),
    ("underscore", _UNDERSCORE_RE),
    ("xrun", _XRUN_RE),
]

_SOURCE_NOTE = {
    "underscore": "fill-in line",
    "bracket": "bracketed blank",
    "checkbox": "checkbox, select if applicable",
    "xrun": "placeholder to replace",
    "label_colon": "value for this label",
    "highlight": "highlighted field, fill from case file",
    "table": "table column, one entry per row",
}


@dataclass
class _Span:
    start: int
    end: int
    source: str
    inner: str  # original matched text (for highlight: the highlighted value)


# --------------------------------------------------------------------------- #
# Key / label / instruction derivation
# --------------------------------------------------------------------------- #
def _humanize(key: str) -> str:
    s = re.sub(r"[_\-.]+", " ", key.strip())
    s = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s.title() if s else key


def _slug(label: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "_", label.strip().lower()).strip("_")
    return s or "blank"


def _label_before(text: str, start: int) -> Optional[str]:
    """Find a 'Label :' or leading 'Label:' to the left of a blank, same line."""
    line = text[max(0, start - 80):start].replace(HL_OPEN, "").replace(HL_CLOSE, "")
    line = line.rsplit("\n", 1)[-1]
    m = re.search(r"([A-Za-z][A-Za-z0-9 ./&'-]{1,40}?)\s*:\s*$", line)
    if m:
        return m.group(1).strip()
    m = re.match(r"\s*([A-Z][A-Za-z0-9 ./&'-]{1,40}?)\s*:", line)
    if m:
        return m.group(1).strip()
    return None


def _derive_key_label(text: str, sp: _Span, idx: int, used: set) -> Tuple[str, str]:
    label = _label_before(text, sp.start)
    if not label:
        inner = sp.inner.strip(" _[]")
        if sp.source == "xrun" and "@" in sp.inner:
            label = "email"
        elif sp.source in ("highlight", "xrun") and re.fullmatch(r"[A-Za-z][A-Za-z ]{1,30}", inner or ""):
            label = inner
    base = _slug(label) if label else f"blank_{idx}"
    key, n = base, 2
    while key in used:
        key = f"{base}_{n}"
        n += 1
    used.add(key)
    return key, _humanize(key)


def _derive_instruction(text: str, sp: _Span) -> str:
    note = _SOURCE_NOTE.get(sp.source, "blank")
    s, e = max(0, sp.start - 40), min(len(text), sp.end + 25)
    ctx = text[s:e].replace(HL_OPEN, "").replace(HL_CLOSE, "")
    instr = f"{note}; context: {ctx}"
    # Strip characters that would corrupt the {{key | instruction}} token.
    instr = re.sub(r"[\[\]{}|_]+", " ", instr)
    return re.sub(r"\s+", " ", instr).strip()


# --------------------------------------------------------------------------- #
# Span finding + overlap / adjacency resolution
# --------------------------------------------------------------------------- #
def _near_markup(sp: _Span, protected: List[Tuple[int, int]]) -> bool:
    """A fill-line that hugs an existing {{markup}} field is already declared."""
    return any(
        (ps - sp.end) in range(0, 31) or (sp.start - pe) in range(0, 31)
        for ps, pe in protected
    )


def find_spans(text: str) -> List[_Span]:
    protected = [(m.start(), m.end()) for m in _MARKUP_RE.finditer(text)]

    def is_protected(a: int, b: int) -> bool:
        return any(a < pe and ps < b for ps, pe in protected)

    cands: List[_Span] = []
    for m in re.finditer(re.escape(HL_OPEN) + r"(.*?)" + re.escape(HL_CLOSE), text, re.S):
        cands.append(_Span(m.start(), m.end(), "highlight", m.group(1)))
    for src, rx in _TEXT_DETECTORS:
        for m in rx.finditer(text):
            cands.append(_Span(m.start(), m.end(), src, m.group(0)))

    # Greedy: earliest start first, longest wins ties; drop protected/overlap.
    cands.sort(key=lambda s: (s.start, -(s.end - s.start)))
    chosen: List[_Span] = []
    occ: List[Tuple[int, int]] = []
    for sp in cands:
        if is_protected(sp.start, sp.end):
            continue
        if any(sp.start < e and s < sp.end for s, e in occ):
            continue
        if sp.source == "underscore" and _near_markup(sp, protected):
            continue
        chosen.append(sp)
        occ.append((sp.start, sp.end))
    chosen.sort(key=lambda s: s.start)
    return chosen


def normalize(text: str) -> str:
    """Rewrite every Tier-2 blank in `text` to canonical {{key | instruction}}.

    Pre-existing {{}} / [[]] markup is preserved verbatim. The result is plain
    text the existing template machinery (detect_fields / fill_text) consumes
    unchanged.
    """
    spans = find_spans(text)
    out: List[str] = []
    used: set = set()
    pos = idx = 0
    for sp in spans:
        if sp.start < pos:
            continue
        out.append(text[pos:sp.start])
        idx += 1
        key, _ = _derive_key_label(text, sp, idx, used)
        instr = _derive_instruction(text, sp)
        out.append("{{" + key + " | " + instr + "}}")
        pos = sp.end
    out.append(text[pos:])
    return "".join(out).replace(HL_OPEN, "").replace(HL_CLOSE, "")


# --------------------------------------------------------------------------- #
# .docx reading at the OOXML run level (FR-5.2): preserve the highlight signal
# --------------------------------------------------------------------------- #
def _paragraph_text(paragraph) -> str:
    """Flatten a paragraph, wrapping contiguous highlighted runs in sentinels."""
    parts: List[str] = []
    buf: List[str] = []
    in_hl = False

    def flush() -> None:
        nonlocal buf
        s = "".join(buf)
        parts.append(HL_OPEN + s + HL_CLOSE if s.strip() else s)
        buf = []

    for run in paragraph.runs:
        try:
            hot = run.font.highlight_color is not None
        except Exception:
            hot = False
        if hot:
            in_hl = True
            buf.append(run.text)
        else:
            if in_hl:
                flush()
                in_hl = False
            parts.append(run.text)
    if in_hl:
        flush()
    out = "".join(parts)
    return out if out else paragraph.text


def _table_lines(table) -> List[str]:
    """Emit a synthetic 'Header: ____' blank for each entirely-empty column of a
    header-bearing table (the empty fill-in grids real forms use)."""
    rows = table.rows
    if len(rows) < 2:
        return []
    headers = [c.text.strip() for c in rows[0].cells]
    body = rows[1:]
    lines: List[str] = []
    for j, header in enumerate(headers):
        if not header:
            continue
        col_filled = any(
            (r.cells[j].text.strip() if j < len(r.cells) else "") for r in body
        )
        if not col_filled:  # entirely empty column => a fill-in field
            lines.append(f"{header}: ____")
    return lines


def read_docx_text(path: str) -> str:
    """Read a .docx into normalized-ready text, preserving highlight + empty
    table grids as blanks. Body order (paragraphs vs. tables) is approximated by
    appending detected table columns after the paragraph stream — sufficient for
    detection and filling in the prototype."""
    from docx import Document  # python-docx

    doc = Document(path)
    parts = [_paragraph_text(p) for p in doc.paragraphs]
    for table in doc.tables:
        parts.extend(_table_lines(table))
    return "\n".join(parts)
