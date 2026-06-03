"""
Template handling (FR-4, FR-5).

Parses templates and detects blanks expressed as {{key}} or [[key]], optionally
with an authoring instruction: {{key | instruction}}. Derives a human-readable
label for each blank.
"""
from __future__ import annotations

import os
import re
from typing import List, Tuple

from .models import FieldSpec

# Matches {{ key }}, {{ key | instruction }}, [[ key ]], [[ key | instruction ]]
_BLANK_RE = re.compile(
    r"(?P<token>(?:\{\{|\[\[)\s*(?P<key>[^|}\]]+?)\s*(?:\|\s*(?P<instr>[^}\]]+?)\s*)?(?:\}\}|\]\]))"
)


def humanize(key: str) -> str:
    """Derive a human-readable label from a field key (FR-5)."""
    s = re.sub(r"[_\-.]+", " ", key.strip())
    s = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", s)  # split camelCase
    s = re.sub(r"\s+", " ", s).strip()
    return s.title() if s else key


def detect_fields(text: str) -> List[FieldSpec]:
    """Detect blanks in order of first appearance, de-duplicated by key."""
    fields: List[FieldSpec] = []
    seen = set()
    for m in _BLANK_RE.finditer(text):
        key = (m.group("key") or "").strip()
        if not key or key in seen:
            continue
        seen.add(key)
        instr = m.group("instr")
        fields.append(
            FieldSpec(
                key=key,
                label=humanize(key),
                instruction=instr.strip() if instr else None,
                placeholder=m.group("token"),
            )
        )
    return fields


def prepare_template(text: str) -> Tuple[str, List[FieldSpec]]:
    """
    Resolve a template to (canonical_text, fields), the single source of truth
    for both the blank count shown to the attorney and the text the filler fills.

    Precedence:
      1. If the template already uses canonical ``{{key}}`` / ``[[key]]`` markup,
         honor it exactly (highest precision, zero guessing).
      2. Otherwise run Tier-2 deterministic detection (``blank_detection``) to
         normalize real firm conventions (underscores, brackets, checkboxes,
         ``Label :`` pairs, sentinels) into canonical ``{{key}}`` markup so the
         template no longer detects zero blanks. (FR-5.1, FR-5.3)
    """
    mustache = detect_fields(text)
    if mustache:
        return text, mustache
    from .blank_detection import normalize

    canonical_text, fields, _candidates = normalize(text)
    return canonical_text, fields


def read_template_text(path: str) -> Tuple[str, str]:
    """Return (kind, text) for a template file. Supports docx / txt / md."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        from docx import Document

        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "docx", "\n".join(parts)
    if ext in (".txt", ".md"):
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            return ext.lstrip("."), fh.read()
    raise ValueError(f"unsupported template type: {ext}")


def fill_text(template_text: str, values: dict) -> str:
    """
    Replace each detected blank token with its filled value. Blanks that map to
    a NEEDS_REVIEW / missing value keep a visible review marker rather than being
    silently emptied (FR-8).
    """
    def _sub(m: re.Match) -> str:
        key = (m.group("key") or "").strip()
        if key in values and values[key] is not None:
            return str(values[key])
        return f"[[NEEDS REVIEW: {humanize(key)}]]"

    return _BLANK_RE.sub(_sub, template_text)
