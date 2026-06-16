"""
Ingestion and chunking (FR-1, FR-2).

Reads .pdf / .docx / .txt / .md / .eml / .xlsx from a per-matter folder and
splits each document into overlapping text chunks suitable for retrieval.
Parsing is best-effort: a single unreadable file never aborts a matter.
"""
from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".eml", ".xlsx"}

# Bound how deep we follow forwarded-email-as-attachment chains, so a
# pathological (or malicious) .eml that nests itself can't recurse forever.
_MAX_EML_DEPTH = 3


@dataclass
class Chunk:
    document: str  # source filename
    text: str
    index: int  # chunk index within the document
    page: Optional[int] = None  # 1-based source page (PDFs); None when not paginated


@dataclass
class IngestedDoc:
    filename: str
    text: str
    chunks: List[Chunk] = field(default_factory=list)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def _read_docx(path: str) -> str:
    from docx import Document  # python-docx

    doc = Document(path)
    parts: List[str] = [p.text for p in doc.paragraphs]
    # include table cell text, which legal docs use heavily
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _strip_html(html: str) -> str:
    """Crude tag strip for the HTML body of an email that has no plain-text part."""
    html = re.sub(r"(?is)<(script|style)\b.*?</\1>", " ", html)
    text = re.sub(r"(?s)<[^>]+>", " ", html)
    # collapse entities we care about and whitespace
    text = (
        text.replace("&nbsp;", " ")
        .replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&#39;", "'")
        .replace("&quot;", '"')
    )
    return re.sub(r"[ \t]*\n[ \t]*", "\n", re.sub(r"[ \t]+", " ", text)).strip()


def _read_eml(path: str, depth: int = 0) -> str:
    """Extract an email's key headers, body text, and attachment content.

    Correspondence is central to a matter (a demand, a settlement offer, an
    engagement). We surface From/To/Date/Subject plus the message body so a fact
    like a date or dollar figure stated in an email is groundable — and we
    recurse into supported attachments (the signed exhibit PDF, the damages
    .xlsx, a forwarded .eml), since that is usually where the operative document
    actually lives. Unsupported attachments (images, zips) are listed by name
    only. The attachment's text is labeled so provenance stays meaningful."""
    import email
    from email import policy

    with open(path, "rb") as fh:
        msg = email.message_from_binary_file(fh, policy=policy.default)

    parts: List[str] = []
    for h in ("From", "To", "Cc", "Date", "Subject"):
        v = msg.get(h)
        if v:
            parts.append(f"{h}: {v}")

    body = ""
    try:
        chosen = msg.get_body(preferencelist=("plain", "html"))
        if chosen is not None:
            content = chosen.get_content()
            if chosen.get_content_type() == "text/html":
                content = _strip_html(content)
            body = content
    except Exception:
        # Malformed MIME: fall back to the first text/plain part we can find.
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                try:
                    body = part.get_content()
                    break
                except Exception:
                    continue
    if body:
        parts.append("")
        parts.append(body.strip())

    names: List[str] = []
    extracted: List[str] = []
    try:
        attachments = list(msg.iter_attachments())
    except Exception:
        attachments = []
    for part in attachments:
        ctype = part.get_content_type()
        fname = part.get_filename() or (
            "(embedded email)" if ctype == "message/rfc822" else f"(unnamed {ctype})"
        )
        names.append(fname)
        if depth >= _MAX_EML_DEPTH:
            continue
        text = _attachment_text(part, fname, ctype, depth)
        if text and text.strip():
            extracted.append(f"--- Attachment: {fname} ---\n{text.strip()}")

    if names:
        parts.append("")
        parts.append("Attachments: " + ", ".join(names))
    for block in extracted:
        parts.append("")
        parts.append(block)

    return "\n".join(parts)


def _attachment_text(part, fname: str, ctype: str, depth: int) -> str:
    """Extract text from a supported email attachment by writing it to a temp
    file and routing through the normal readers — so an attached PDF even gets
    the OCR fallback, and a forwarded .eml recurses (with the depth guard).
    Unsupported types return '' and surface as a name only."""
    import tempfile

    ext = os.path.splitext(fname)[1].lower()
    if ctype == "message/rfc822":
        ext = ".eml"
    if ext not in SUPPORTED_EXTS:
        return ""

    data = part.get_payload(decode=True)
    if data is None and ctype == "message/rfc822":
        # message/rfc822 parts aren't base64 leaves; serialize the embedded message.
        payload = part.get_payload()
        if isinstance(payload, list) and payload:
            try:
                data = payload[0].as_bytes()
            except Exception:
                data = None
    if not data:
        return ""

    fd, tmp = tempfile.mkstemp(suffix=ext)
    try:
        with os.fdopen(fd, "wb") as tf:
            tf.write(data)
        if ext == ".eml":
            return _read_eml(tmp, depth=depth + 1)
        return read_document(tmp)
    except Exception as exc:
        print(f"[ingest] could not read attachment {fname}: {exc}")
        return ""
    finally:
        try:
            os.remove(tmp)
        except OSError:
            pass


def _read_xlsx(path: str) -> str:
    """Flatten a spreadsheet to text, sheet by sheet, tab-separated rows.

    Legal matters carry damages calculations, billing/expense ledgers, and asset
    schedules as .xlsx. We read cached cell *values* (data_only) so a figure
    produced by a formula is transcribable, and label each sheet so provenance
    stays meaningful."""
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    parts: List[str] = []
    try:
        for ws in wb.worksheets:
            parts.append(f"[Sheet: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(c.strip() for c in cells):
                    parts.append("\t".join(cells))
    finally:
        wb.close()
    return "\n".join(parts)


def _read_pdf(path: str) -> str:
    text = ""
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        if reader.is_encrypted:
            # Many e-filed / "signed" PDFs are encrypted with an empty user
            # password; decrypt needs the `cryptography` package for AES.
            try:
                reader.decrypt("")
            except Exception:
                pass
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:
        # Encrypted-without-cryptography, password-protected, or corrupt: don't
        # abandon the file — fall through to OCR, which can often still render it.
        print(f"[ingest] pypdf could not extract {os.path.basename(path)} ({exc}); trying OCR.")
    # Scanned/image (or unreadable) PDFs have no usable text layer — try OCR (FR-1).
    return _with_ocr_fallback(text, path)


# --------------------------------------------------------------------------- #
# OCR fallback for scanned PDFs (FR-1). Optional + graceful: if Tesseract /
# poppler aren't installed, extraction simply returns empty and the fill's
# diagnostic flags the document — nothing crashes (NFR-3, NFR-5).
# --------------------------------------------------------------------------- #
OCR_MIN_CHARS = 40  # below this a PDF is treated as scanned and OCR is attempted
OCR_DPI = int(os.environ.get("VERBATIM_OCR_DPI", "200"))
_OCR_ENABLED = os.environ.get("VERBATIM_OCR", "1") != "0"
# Optional explicit binary locations, so a Windows host needn't edit PATH:
#   VERBATIM_TESSERACT_CMD = C:\Program Files\Tesseract-OCR\tesseract.exe
#   VERBATIM_POPPLER_PATH  = C:\poppler-24.08.0\Library\bin
_TESSERACT_CMD = os.environ.get("VERBATIM_TESSERACT_CMD")
_POPPLER_PATH = os.environ.get("VERBATIM_POPPLER_PATH")
# Where eng.traineddata lives, if not alongside tesseract.exe (e.g. a separate
# tessdata folder): VERBATIM_TESSDATA_DIR = D:\Program Files\tessdata
_TESSDATA_DIR = os.environ.get("VERBATIM_TESSDATA_DIR")


def _configure_tesseract(pytesseract) -> None:
    if _TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = _TESSERACT_CMD
    if _TESSDATA_DIR:
        os.environ["TESSDATA_PREFIX"] = _TESSDATA_DIR


def _tesseract_ok() -> bool:
    try:
        import pytesseract
        from pdf2image import convert_from_path  # noqa: F401

        _configure_tesseract(pytesseract)
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _poppler_ok() -> bool:
    """Verify poppler's pdftoppm is reachable — OCR rasterization needs it, and
    its absence is the most common reason OCR silently produces nothing."""
    import shutil

    exe = "pdftoppm"
    if _POPPLER_PATH:
        win = os.path.join(_POPPLER_PATH, exe + ".exe")
        nix = os.path.join(_POPPLER_PATH, exe)
        if os.path.isfile(win) or os.path.isfile(nix):
            return True
        return shutil.which(exe, path=_POPPLER_PATH) is not None
    return shutil.which(exe) is not None or shutil.which("pdfinfo") is not None


def ocr_available() -> bool:
    """True only if OCR can actually run: enabled, Tesseract, AND poppler."""
    return _OCR_ENABLED and _tesseract_ok() and _poppler_ok()


def _ocr_pdf_pages(path: str) -> List[str]:
    """OCR a PDF, returning one text string per page ('' if OCR isn't available)."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except Exception as exc:
        print(f"[ingest] OCR libraries unavailable ({exc}); install pytesseract + pdf2image.")
        return []
    _configure_tesseract(pytesseract)
    kwargs = {"dpi": OCR_DPI}
    if _POPPLER_PATH:
        kwargs["poppler_path"] = _POPPLER_PATH
    try:
        images = convert_from_path(path, **kwargs)
    except Exception as exc:
        print(
            f"[ingest] could not rasterize {path} for OCR ({exc}); is poppler installed / "
            f"VERBATIM_POPPLER_PATH set to its bin folder?"
        )
        return []
    pages: List[str] = []
    for img in images:
        try:
            pages.append(pytesseract.image_to_string(img))
        except Exception as exc:
            print(f"[ingest] OCR failed on a page of {path}: {exc}")
            pages.append("")
    return pages


def ocr_pdf(path: str) -> str:
    """Rasterize a PDF and OCR each page, joined. Returns '' if OCR isn't available."""
    return "\n".join(_ocr_pdf_pages(path)).strip()


def _with_ocr_fallback(text: str, path: str) -> str:
    """Use OCR text when the native PDF text layer is empty/negligible."""
    if len(text.strip()) >= OCR_MIN_CHARS or not _OCR_ENABLED:
        return text
    ocr = ocr_pdf(path)
    return ocr if len(ocr.strip()) > len(text.strip()) else text


def _pdf_pages(path: str) -> List[str]:
    """Per-page text of a PDF (so provenance can cite a page), with per-page OCR
    fallback for scanned/encrypted files."""
    texts: List[str] = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                pass
        texts = [(p.extract_text() or "") for p in reader.pages]
    except Exception as exc:
        print(f"[ingest] pypdf could not extract {os.path.basename(path)} ({exc}); trying OCR.")
    if sum(len(t.strip()) for t in texts) < OCR_MIN_CHARS and _OCR_ENABLED:
        ocr = _ocr_pdf_pages(path)
        if any(t.strip() for t in ocr):
            return ocr
    return texts


def read_document(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext in (".txt", ".md"):
        return _read_txt(path)
    if ext == ".eml":
        return _read_eml(path)
    if ext == ".xlsx":
        return _read_xlsx(path)
    raise ValueError(f"unsupported extension: {ext}")


def read_document_pages(path: str) -> List[Tuple[Optional[int], str]]:
    """Return [(page, text), ...]. PDFs yield one entry per 1-based page so a
    filled value can cite the page it came from; other formats are unpaginated
    and yield a single (None, text) entry."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return [(i + 1, t) for i, t in enumerate(_pdf_pages(path))]
    if ext == ".docx":
        return [(None, _read_docx(path))]
    if ext in (".txt", ".md"):
        return [(None, _read_txt(path))]
    if ext == ".eml":
        return [(None, _read_eml(path))]
    if ext == ".xlsx":
        return [(None, _read_xlsx(path))]
    raise ValueError(f"unsupported extension: {ext}")


def chunk_text(
    text: str, document: str, size: int = 1100, overlap: int = 200, page: Optional[int] = None
) -> List[Chunk]:
    """Split into overlapping character windows, preferring paragraph breaks."""
    text = text.strip()
    if not text:
        return []
    chunks: List[Chunk] = []
    start = 0
    idx = 0
    n = len(text)
    while start < n:
        end = min(start + size, n)
        # try to end on a paragraph or sentence boundary for cleaner passages
        if end < n:
            window = text[start:end]
            for sep in ("\n\n", "\n", ". "):
                cut = window.rfind(sep)
                if cut > size * 0.5:
                    end = start + cut + len(sep)
                    break
        chunks.append(Chunk(document=document, text=text[start:end].strip(), index=idx, page=page))
        idx += 1
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return [c for c in chunks if c.text]


def ingest_folder(folder: str) -> List[IngestedDoc]:
    """Ingest every supported document in a matter folder (non-recursive + 1 level)."""
    docs: List[IngestedDoc] = []
    if not os.path.isdir(folder):
        return docs
    for root, _dirs, files in os.walk(folder):
        for name in sorted(files):
            ext = os.path.splitext(name)[1].lower()
            if ext not in SUPPORTED_EXTS:
                continue
            path = os.path.join(root, name)
            rel = os.path.relpath(path, folder)
            try:
                pages = read_document_pages(path)
            except Exception as exc:  # best-effort; skip unreadable files
                pages = []
                print(f"[ingest] could not read {name}: {exc}")
            chunks: List[Chunk] = []
            for pg, ptext in pages:
                chunks.extend(chunk_text(ptext, rel, page=pg))
            full_text = "\n".join(t for _pg, t in pages)
            docs.append(IngestedDoc(filename=rel, text=full_text, chunks=chunks))
    return docs


def total_chars(docs: List[IngestedDoc]) -> int:
    return sum(len(d.text) for d in docs)
